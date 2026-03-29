from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import arabic_reshaper
from bidi.algorithm import get_display


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(parents=True, exist_ok=True)


CONTENT = {
    "cover": {
        "title": "Echorouk Editorial OS",
        "subtitle_ar": "منصة غرفة الأخبار الذكية",
        "subtitle_fr": "Plateforme de Rédaction Intelligente",
        "subtitle_en": "Intelligent Newsroom Platform",
        "audience": "غرف الأخبار | Newsrooms | Rédactions",
        "version": "الإصدار: 1.0",
        "date": "التاريخ: 29 مارس 2026",
        "image_note": "صورة الواجهة توضع هنا",
    },
    "toc": [
        "1. تعريف مختصر بالمنصة",
        "2. لماذا هذه المنصة؟",
        "3. كيف تعمل داخل غرفة الأخبار",
        "4. الوحدات والخصائص الأساسية",
        "5. الأمان والحوكمة والتحكم",
        "6. التكامل والتشغيل",
        "7. مؤشرات النجاح والدعم",
    ],
    "ar": {
        "title": "الوثيقة التعريفية الرسمية – العربية",
        "sections": [
            ("1) تعريف مختصر بالمنصة", [
                "منصة غرفة الأخبار الذكية هي نظام عمل متكامل يدير دورة الخبر بالكامل، من الاستقبال والتحرير وحتى الاعتماد والنشر.",
                "تجمع المنصة بين أدوات تحرير ذكية، تقارير تدقيق وتحقق، ومسار اعتماد واضح يضمن الجودة قبل النشر.",
            ]),
            ("2) لماذا هذه المنصة؟", [
                "توحيد مسار العمل داخل غرفة الأخبار وتقليل التشتيت بين الأدوات.",
                "تسريع دورة الإنتاج مع الحفاظ على جودة النص ومصداقيته.",
                "رفع جاهزية المقال للنشر الرقمي من خلال SEO والسوشيال.",
                "تقليل الأخطاء التحريرية عبر التقارير الآلية قبل الاعتماد.",
            ]),
            ("3) كيف تعمل داخل غرفة الأخبار", [
                "استقبال الأخبار تلقائيًا من المصادر أو من الأرشيف.",
                "تحويلها إلى مسودة قابلة للتحرير داخل المحرر الذكي.",
                "تشغيل أدوات التدقيق والتأكد من الادعاءات وتحسين السيو.",
                "إرسال المادة للاعتماد أو إعادة المراجعة بشكل واضح.",
                "تجهيز نسخة نهائية للنشر أو للنشر الرقمي المتعدد.",
            ]),
            ("4) الوحدات والخصائص الأساسية", [
                "محرر ذكي متعدد الأدوات (تدقيق، تحقق، تحسين، نسخ سوشيال).",
                "تقارير جودة تحريرية قبل النشر (بوابات جودة).",
                "إدارة الأرشيف والنسخ والتعديلات التاريخية.",
                "تغطية رقمية جاهزة لمنصات السوشيال والإشعارات.",
                "SEO متوافق مع شروط Yoast وتحسينات تلقائية.",
            ]),
            ("5) الأمان والحوكمة والتحكم", [
                "توزيع الأدوار حسب الصلاحيات (صحفي، رئيس تحرير، مدير).",
                "مسار اعتماد واضح قابل للتتبع والتدقيق.",
                "سجل تغييرات كامل لكل مادة لضمان الحوكمة.",
            ]),
            ("6) التكامل والتشغيل", [
                "تشغيل داخلي أو سحابي بحسب سياسة المؤسسة.",
                "تكامل مع مصادر RSS أو أنظمة أرشفة داخلية.",
                "قابل للتخصيص وفق معايير التحرير المحلية.",
            ]),
            ("7) مؤشرات النجاح والدعم", [
                "تقليل وقت تجهيز المادة للنشر.",
                "رفع جودة النص النهائي قبل الاعتماد.",
                "دعم فني مباشر وتحديثات دورية.",
            ]),
        ],
    },
    "fr": {
        "title": "Document de Présentation – Français",
        "sections": [
            ("1) Présentation rapide", [
                "Editorial OS est un système complet qui gère le cycle de vie d’un article, de la réception jusqu’à la publication.",
                "Il combine un éditeur intelligent, des rapports de qualité et un workflow clair d’approbation.",
            ]),
            ("2) Pourquoi cette plateforme ?", [
                "Unifie le travail rédactionnel et réduit la fragmentation.",
                "Accélère la production tout en assurant la qualité.",
                "Améliore la performance SEO et social media.",
                "Réduit les erreurs grâce à des contrôles automatiques.",
            ]),
            ("3) Fonctionnement en rédaction", [
                "Collecte automatique des actualités.",
                "Transformation en brouillons éditables.",
                "Activation des outils (vérification, SEO, social).",
                "Validation par les responsables éditoriaux.",
                "Préparation finale pour la publication multi‑canal.",
            ]),
            ("4) Modules essentiels", [
                "Éditeur intelligent multi‑outils.",
                "Rapports de qualité avant publication.",
                "Gestion d’archives et versions.",
                "Pack digital pour réseaux sociaux et notifications.",
                "SEO optimisé selon Yoast.",
            ]),
            ("5) Sécurité & gouvernance", [
                "Rôles et permissions distincts.",
                "Workflow d’approbation traçable.",
                "Historique complet des modifications.",
            ]),
            ("6) Intégration & déploiement", [
                "Déploiement cloud ou on‑premise.",
                "Compatibilité RSS et archives internes.",
                "Personnalisation selon les règles éditoriales.",
            ]),
            ("7) Résultats & support", [
                "Gain de temps dans la publication.",
                "Qualité éditoriale renforcée.",
                "Support technique continu.",
            ]),
        ],
    },
    "en": {
        "title": "Official Client Presentation – English",
        "sections": [
            ("1) Quick Overview", [
                "Editorial OS is a full newsroom platform that manages the entire news lifecycle.",
                "It combines a smart editor, quality reports, and a clear approval workflow.",
            ]),
            ("2) Why This Platform?", [
                "Unifies newsroom workflows and reduces tool fragmentation.",
                "Speeds up production while improving quality.",
                "Boosts SEO and social readiness.",
                "Reduces editorial errors via automated checks.",
            ]),
            ("3) How It Works", [
                "Automatic news intake and sourcing.",
                "Editable drafts in the smart editor.",
                "Run verification, SEO, and social tools.",
                "Approval by editors and chiefs.",
                "Final preparation for multi‑channel publishing.",
            ]),
            ("4) Core Modules", [
                "Smart multi‑tool editor.",
                "Quality and verification gates.",
                "Archive & version history.",
                "Digital publishing pack for social and push.",
                "Yoast‑compatible SEO generation.",
            ]),
            ("5) Security & Governance", [
                "Role‑based permissions.",
                "Trackable approval chain.",
                "Full editorial audit trail.",
            ]),
            ("6) Integration & Deployment", [
                "On‑premise or cloud deployment.",
                "RSS and internal archive integration.",
                "Customizable to editorial policies.",
            ]),
            ("7) Success Metrics & Support", [
                "Faster publishing cycle.",
                "Higher editorial quality.",
                "Dedicated support and updates.",
            ]),
        ],
    },
}


def set_rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading(doc, text, level=1, rtl=False):
    p = doc.add_heading(text, level=level)
    if rtl:
        set_rtl(p)
    return p


def add_paragraph(doc, text, rtl=False, style=None):
    p = doc.add_paragraph(text, style=style)
    if rtl:
        set_rtl(p)
    return p


def build_docx(path: Path):
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"

    cover = CONTENT["cover"]
    add_heading(doc, cover["title"], level=0)
    add_paragraph(doc, cover["subtitle_ar"], rtl=True)
    add_paragraph(doc, cover["subtitle_fr"])
    add_paragraph(doc, cover["subtitle_en"])
    add_paragraph(doc, cover["audience"], rtl=True)
    add_paragraph(doc, cover["version"], rtl=True)
    add_paragraph(doc, cover["date"], rtl=True)
    add_paragraph(doc, f"[{cover['image_note']}]", rtl=True)

    doc.add_page_break()

    add_heading(doc, "الفهرس | Table of Contents | Sommaire", level=1, rtl=True)
    for item in CONTENT["toc"]:
        add_paragraph(doc, item, rtl=True, style="List Number")

    doc.add_page_break()

    add_heading(doc, CONTENT["ar"]["title"], level=1, rtl=True)
    for title, paragraphs in CONTENT["ar"]["sections"]:
        add_heading(doc, title, level=2, rtl=True)
        for para in paragraphs:
            add_paragraph(doc, para, rtl=True)

    doc.add_page_break()

    add_heading(doc, CONTENT["fr"]["title"], level=1)
    for title, paragraphs in CONTENT["fr"]["sections"]:
        add_heading(doc, title, level=2)
        for para in paragraphs:
            add_paragraph(doc, para)

    doc.add_page_break()

    add_heading(doc, CONTENT["en"]["title"], level=1)
    for title, paragraphs in CONTENT["en"]["sections"]:
        add_heading(doc, title, level=2)
        for para in paragraphs:
            add_paragraph(doc, para)

    doc.save(path)


def wrap_text(text, max_width, font_name, font_size):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        width = pdfmetrics.stringWidth(candidate, font_name, font_size)
        if width <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def wrap_text_rtl(text, max_width, font_name, font_size):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        reshaped = get_display(arabic_reshaper.reshape(candidate))
        width = pdfmetrics.stringWidth(reshaped, font_name, font_size)
        if width <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_paragraph(c, text, x, y, max_width, font_name, font_size, leading, rtl=False):
    if rtl:
        lines = wrap_text_rtl(text, max_width, font_name, font_size)
        for line in lines:
            shaped = get_display(arabic_reshaper.reshape(line))
            width = pdfmetrics.stringWidth(shaped, font_name, font_size)
            c.drawString(x + max_width - width, y, shaped)
            y -= leading
        return y
    lines = wrap_text(text, max_width, font_name, font_size)
    for line in lines:
        c.drawString(x, y, line)
        y -= leading
    return y


def build_pdf(path: Path):
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    pdfmetrics.registerFont(TTFont("Arial", str(font_path)))

    width, height = A4
    c = canvas.Canvas(str(path), pagesize=A4)
    margin = 2 * cm
    max_width = width - 2 * margin

    # Cover page
    cover = CONTENT["cover"]
    c.setFont("Arial", 24)
    c.drawString(margin, height - 3 * cm, cover["title"])
    c.setFont("Arial", 14)
    draw_paragraph(c, cover["subtitle_ar"], margin, height - 4.5 * cm, max_width, "Arial", 14, 16, rtl=True)
    draw_paragraph(c, cover["subtitle_fr"], margin, height - 5.5 * cm, max_width, "Arial", 12, 15)
    draw_paragraph(c, cover["subtitle_en"], margin, height - 6.2 * cm, max_width, "Arial", 12, 15)
    draw_paragraph(c, cover["audience"], margin, height - 7.2 * cm, max_width, "Arial", 12, 15, rtl=True)
    draw_paragraph(c, cover["version"], margin, height - 8.0 * cm, max_width, "Arial", 11, 14, rtl=True)
    draw_paragraph(c, cover["date"], margin, height - 8.6 * cm, max_width, "Arial", 11, 14, rtl=True)

    c.setStrokeColorRGB(0.6, 0.6, 0.6)
    c.rect(margin, height - 15 * cm, max_width, 5 * cm)
    c.setFont("Arial", 10)
    draw_paragraph(c, cover["image_note"], margin, height - 15.5 * cm, max_width, "Arial", 10, 12, rtl=True)

    # Table of contents
    c.setFont("Arial", 13)
    draw_paragraph(c, "الفهرس | Table of Contents | Sommaire", margin, height - 17 * cm, max_width, "Arial", 13, 16, rtl=True)
    y = height - 18 * cm
    c.setFont("Arial", 11)
    for item in CONTENT["toc"]:
        y = draw_paragraph(c, item, margin, y, max_width, "Arial", 11, 14, rtl=True)

    c.showPage()

    # Arabic page
    c.setFont("Arial", 16)
    y = height - 2 * cm
    y = draw_paragraph(c, CONTENT["ar"]["title"], margin, y, max_width, "Arial", 16, 20, rtl=True)
    c.setFont("Arial", 11)
    for title, paragraphs in CONTENT["ar"]["sections"]:
        y = draw_paragraph(c, title, margin, y - 0.3 * cm, max_width, "Arial", 12, 16, rtl=True)
        for para in paragraphs:
            y = draw_paragraph(c, para, margin, y, max_width, "Arial", 11, 15, rtl=True)
        y -= 0.2 * cm
        if y < 3 * cm:
            c.showPage()
            y = height - 2 * cm
            c.setFont("Arial", 11)

    c.showPage()

    # French page
    c.setFont("Arial", 16)
    y = height - 2 * cm
    y = draw_paragraph(c, CONTENT["fr"]["title"], margin, y, max_width, "Arial", 16, 20)
    c.setFont("Arial", 11)
    for title, paragraphs in CONTENT["fr"]["sections"]:
        y = draw_paragraph(c, title, margin, y - 0.3 * cm, max_width, "Arial", 12, 16)
        for para in paragraphs:
            y = draw_paragraph(c, para, margin, y, max_width, "Arial", 11, 15)
        y -= 0.2 * cm
        if y < 3 * cm:
            c.showPage()
            y = height - 2 * cm
            c.setFont("Arial", 11)

    c.showPage()

    # English page
    c.setFont("Arial", 16)
    y = height - 2 * cm
    y = draw_paragraph(c, CONTENT["en"]["title"], margin, y, max_width, "Arial", 16, 20)
    c.setFont("Arial", 11)
    for title, paragraphs in CONTENT["en"]["sections"]:
        y = draw_paragraph(c, title, margin, y - 0.3 * cm, max_width, "Arial", 12, 16)
        for para in paragraphs:
            y = draw_paragraph(c, para, margin, y, max_width, "Arial", 11, 15)
        y -= 0.2 * cm
        if y < 3 * cm:
            c.showPage()
            y = height - 2 * cm
            c.setFont("Arial", 11)

    c.save()


def main():
    docx_path = OUT_DIR / "Echorouk_Editorial_OS_Client_Presentation.docx"
    pdf_path = OUT_DIR / "Echorouk_Editorial_OS_Client_Presentation.pdf"
    build_docx(docx_path)
    build_pdf(pdf_path)
    print(f"Saved: {docx_path}")
    print(f"Saved: {pdf_path}")


if __name__ == "__main__":
    main()
